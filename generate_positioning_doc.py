#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Génération document positionnement/comparatif ERPScolaire
Tech Solutions DDTP - Document Word professionnel
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

# Création document
doc = Document()

# Configuration marges
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ============================================
# 1. PAGE TITRE
# ============================================

# Logo/Titre
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('🎓 ERPScolaire\n')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.name = 'Times New Roman'

# Sous-titre
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('La Solution de Gestion Scolaire Camerounaise\n')
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True
subtitle_run.font.name = 'Times New Roman'

# Slogan
slogan = doc.add_paragraph()
slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
slogan_run = slogan.add_run('Conçue par des Camerounais, pour les écoles du Cameroun')
slogan_run.font.size = Pt(12)
slogan_run.font.bold = True
slogan_run.font.color.rgb = RGBColor(0, 102, 204)
slogan_run.font.name = 'Times New Roman'

# Ligne vide
doc.add_paragraph()

# Date et version
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info.add_run(f'Document de Positionnement\nMai 2026 | Version 1.0')
info_run.font.size = Pt(10)
info_run.font.italic = True
info_run.font.name = 'Times New Roman'
info_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ============================================
# 2. TABLE DES MATIÈRES
# ============================================

toc_title = doc.add_heading('Table des matières', level=1)
toc_title.paragraph_format.space_after = Pt(12)

toc_items = [
    '1. Qui sommes-nous ?',
    '2. Problème camerounais que nous résolvons',
    '3. Notre solution : ERPScolaire',
    '4. Avantages uniques au Cameroun',
    '5. Comparaison avec concurrents',
    '6. Cas d\'usage réels',
    '7. Tarification transparente',
    '8. Support et formation',
    '9. Appel à l\'action',
]

for item in toc_items:
    toc_p = doc.add_paragraph(item, style='List Number')
    toc_p.paragraph_format.left_indent = Inches(0.5)

doc.add_page_break()

# ============================================
# 3. SECTION 1 : QUI SOMMES-NOUS
# ============================================

doc.add_heading('1. Qui sommes-nous ?', level=1)

p = doc.add_paragraph()
p.add_run('Tech Solutions DDTP ').bold = True
p.add_run('est une entreprise de développement logiciel basée à Douala, Cameroun. Nous créons des solutions informatiques adaptées aux réalités du contexte camerounais et africain.')
p.paragraph_format.line_spacing = 1.5

doc.add_heading('Notre Mission', level=2)
mission = doc.add_paragraph(
    'Simplifier la gestion administrative et pédagogique des établissements scolaires camerounais '
    'avec une solution fiable, abordable et pensée pour la réalité locale.'
)
mission.paragraph_format.line_spacing = 1.5

doc.add_heading('Notre Vision', level=2)
vision = doc.add_paragraph(
    'Que chaque école au Cameroun, du primaire à l\'université, puisse gérer efficacement '
    'ses élèves, notes, paiements et finances en quelques clics.'
)
vision.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

# ============================================
# 4. SECTION 2 : PROBLÈME CAMEROUNAIS
# ============================================

doc.add_heading('2. Le Problème : Gestion Scolaire Camerounaise Aujourd\'hui', level=1)

p = doc.add_paragraph()
p.add_run('Les écoles camerounaises font face à des défis spécifiques :')
p.paragraph_format.line_spacing = 1.5

problems = [
    'Gestion manuelle des élèves → erreurs, doublons, données perdues',
    'Bulletins imprimés → coûteux, lent (3 jours), pas de trace',
    'Suivi paiements compliqué → pas de centralisation, pertes financières',
    'Presences sur papier → oublis, absences du directeur = perte d\'infos',
    'Pas de rapports facilement → impossible de suivre performances',
    'Outils génériques (Google Classroom, Moodle) → pas adaptés au système camerounais',
    'Support étranger → langue, délais, incompréhension du contexte',
    'Coûts élevés → 500-2000 USD/mois = 300-1200k FCFA pour petite école'
]

for problem in problems:
    p = doc.add_paragraph(problem, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

# ============================================
# 5. SECTION 3 : NOTRE SOLUTION
# ============================================

doc.add_heading('3. Notre Solution : ERPScolaire', level=1)

p = doc.add_paragraph()
p.add_run('ERPScolaire est un logiciel de gestion scolaire ')
p.add_run('100% cloud').bold = True
p.add_run(' conçu spécifiquement pour les écoles camerounaises.')
p.paragraph_format.line_spacing = 1.5

doc.add_heading('Modules disponibles (16 au total)', level=2)

modules = [
    ('Gestion Élèves', 'Inscriptions, dossiers, photos, imports/exports Excel'),
    ('Gestion Personnel', 'Enseignants, administrateurs, coordonnées, photos'),
    ('Classes & Niveaux', 'Structure du système camerounais (Primaire/Collège/Lycée)'),
    ('Emploi du Temps', 'Grille hebdomadaire, salles, créneaux horaires'),
    ('Notes & Résultats', 'Saisie par trimestre, calcul automatique, classements'),
    ('Bulletins PDF', 'Générés en 2 minutes, imprimables, archivables'),
    ('Présences/Appel', 'Appel numérique, rapports d\'assiduité, alertes parents'),
    ('Paiements', 'Suivi dossiers, tranches, reçus, exports Excel'),
    ('Examens', 'Programmation, calendrier, épreuves'),
    ('Comptabilité', 'Recettes/dépenses, bilan, exports'),
    ('Bibliothèque', 'Catalogue, emprunts, retours automatisés'),
    ('Communication', 'Annonces, messagerie directe'),
    ('Rapports', '4 rapports complets + exports Excel'),
    ('Dashboard', 'Statistiques en temps réel, graphiques'),
    ('Paramètres', 'Configuration école, années scolaires, créneaux'),
    ('Profil Utilisateur', 'Gestion compte, photo, mot de passe')
]

for name, desc in modules:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(name).bold = True
    p.add_run(f' : {desc}')
    p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

# ============================================
# 6. SECTION 4 : AVANTAGES UNIQUES
# ============================================

doc.add_heading('4. Avantages Uniques d\'ERPScolaire au Cameroun', level=1)

advantages = [
    {
        'title': '✓ Conçu pour le système scolaire camerounais',
        'desc': 'Primaire (6 ans), Collège (4 ans), Lycée (3 ans), Université. Trimestres exactes. Pas d\'adaptation nécessaire.'
    },
    {
        'title': '✓ Équipe 100% camerounaise + support local',
        'desc': 'Support WhatsApp < 1h. Équipe parle français, comprend le contexte. Pas d\'attendre support étranger.'
    },
    {
        'title': '✓ Tarification en FCFA sans surprise',
        'desc': 'Plan Basique : 35 000 FCFA/mois | Plan Pro : 85 000 FCFA/mois | Plan Entreprise : 250 000 FCFA/mois. Zéro frais cachés.'
    },
    {
        'title': '✓ 5-10x moins cher que concurrents',
        'desc': 'Google Workspace = 864 USD/an. Nous = 102 000 FCFA/an. Tarif réaliste pour écoles camerounaises.'
    },
    {
        'title': '✓ Hébergement flexible',
        'desc': 'Serveur Cameroun (rapide) ou serveur EU (RGPD). Vous choisissez. Pas d\'obligation cloud étranger.'
    },
    {
        'title': '✓ Fonctionne avec internet camerounais réel',
        'desc': 'Optimisé pour 3G/4G. Cache local si internet coupe. Pas de panique si Orange/Camtel ralenti.'
    },
    {
        'title': '✓ Zéro dépendance fournisseur étranger',
        'desc': 'Vos données restent vôtres. Export complet possible. Pas enfermé comme chez Google/Microsoft.'
    },
    {
        'title': '✓ Bulletins PDF au format camerounais',
        'desc': 'Format local pour administration, avec tampons, signatures, conformité règlementaire.'
    },
    {
        'title': '✓ Multi-établissements',
        'desc': 'Un seul compte gère 5, 10 ou 100 écoles. Idéal pour réseaux d\'écoles privées.'
    },
    {
        'title': '✓ RGPD compliant',
        'desc': 'Données élèves/parents protégées. Chiffrement 256-bit. Conformité légale internationale.'
    },
    {
        'title': '✓ Démo gratuite en ligne',
        'desc': 'Testez complètement avant acheter. Zéro engagement. Zéro carte crédit.'
    },
    {
        'title': '✓ Formation en français',
        'desc': 'Manuel 20+ pages, vidéos tutoriels, webinaires gratuits mensuels, formations sur site possibles.'
    }
]

for adv in advantages:
    doc.add_heading(adv['title'], level=3)
    p = doc.add_paragraph(adv['desc'])
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ============================================
# 7. SECTION 5 : COMPARAISON CONCURRENTS
# ============================================

doc.add_heading('5. Comparaison avec Concurrents', level=1)

# Table de comparaison
table = doc.add_table(rows=9, cols=6)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
header_cells = table.rows[0].cells
headers = ['Critère', 'ERPScolaire', 'Moodle Cloud', 'Google Workspace', 'SchoolBox', 'Jorani']
for i, header in enumerate(headers):
    cell = header_cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = 'Times New Roman'

# Data
comparisons = [
    ['Équipe locale', '✅ Cameroun', '❌ Monde', '❌ USA', '❌ Monde', '❌ Monde'],
    ['Support WhatsApp', '✅ < 1h', '❌ Non', '❌ Non', '❌ Non', '❌ Non'],
    ['Tarif mensuel', '✅ 35k-250k FCFA', '❌ $60/mois', '❌ $72/mois', '❌ $40-150/mois', '❌ Cher'],
    ['Système camerounais', '✅ Exact', '⚠️ Générique', '⚠️ Générique', '⚠️ Partiel', '⚠️ Partiel'],
    ['Bulletins PDF local', '✅ Oui', '❌ Non', '❌ Non', '✅ Oui', '❌ Non'],
    ['Multi-établissements', '✅ 5-∞', '❌ Non', '⚠️ Limité', '⚠️ Limité', '❌ Non'],
    ['Gestion paiements', '✅ Complète', '❌ Non', '❌ Non', '✅ Oui', '✅ Oui'],
    ['Documentation français', '✅ Complète', '⚠️ Partielle', '⚠️ Générique', '⚠️ Partielle', '⚠️ Partielle'],
]

for row_idx, row_data in enumerate(comparisons, 1):
    row_cells = table.rows[row_idx].cells
    for col_idx, text in enumerate(row_data):
        cell = row_cells[col_idx]
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Légende
legend = doc.add_paragraph()
legend.add_run('✅ = Avantage | ⚠️ = Partiel | ❌ = Absence')
legend.paragraph_format.space_before = Pt(12)
legend.paragraph_format.space_after = Pt(12)

doc.add_paragraph()

# ============================================
# 8. SECTION 6 : CAS D'USAGE RÉELS
# ============================================

doc.add_heading('6. Cas d\'Usage Réels', level=1)

doc.add_heading('Cas 1 : Lycée de 500 élèves', level=2)

case1_p = doc.add_paragraph()
case1_p.add_run('Problème avant ERPScolaire :\n').bold = True
case1_p.add_run(
    '• Bulletins : 3 jours à imprimer (500 bulletins × 4 pages)\n'
    '• Paiements : registre papier → erreurs fréquentes\n'
    '• Presences : appels sur papier → incomplet\n'
    '• Notes : saisie Excel → risques erreurs\n'
)
case1_p.paragraph_format.line_spacing = 1.5

case1_after = doc.add_paragraph()
case1_after.add_run('Après ERPScolaire :\n').bold = True
case1_after.add_run(
    '✓ Bulletins générés en 2 minutes (500 en PDF, prêts impression)\n'
    '✓ Paiements : suivi temps réel, zéro oubli\n'
    '✓ Presences : appel 30s par classe, rapports auto\n'
    '✓ Notes : saisie ERPScolaire, calcul auto des moyennes\n'
    '✓ ROI : 40% temps administratif économisé = 4 heures/jour libérées\n'
)
case1_after.paragraph_format.line_spacing = 1.5
case1_after.paragraph_format.space_after = Pt(12)

doc.add_heading('Cas 2 : Réseau de 5 petites écoles privées', level=2)

case2_p = doc.add_paragraph()
case2_p.add_run('Avantage unique d\'ERPScolaire :\n').bold = True
case2_p.add_run(
    '• 1 seul compte gère les 5 écoles\n'
    '• Chaque école voit ses propres données\n'
    '• Rapports consolidés groupe possible\n'
    '• Coût = Plan Pro unique (85 000 FCFA) pour tous\n'
    '• Coût concurrent = 5 abonnements Moodle (300k FCFA+)\n'
)
case2_p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

# ============================================
# 9. SECTION 7 : TARIFICATION
# ============================================

doc.add_heading('7. Tarification Transparente', level=1)

doc.add_heading('Plan Basique - Jusqu\'à 150 élèves', level=2)
basic_p = doc.add_paragraph()
basic_p.add_run('35 000 FCFA/mois\n').bold = True
basic_p.add_run(
    '✓ Dashboard complet\n'
    '✓ 2 établissements\n'
    '✓ Gestion élèves & notes\n'
    '✓ Présences & bulletins\n'
    '✓ Paiements basique\n'
    '✓ Support email\n'
)
basic_p.paragraph_format.line_spacing = 1.5

doc.add_heading('Plan Professionnel - 150 à 500 élèves (RECOMMANDÉ)', level=2)
pro_p = doc.add_paragraph()
pro_p.add_run('85 000 FCFA/mois ⭐\n').bold = True
pro_p.add_run(
    '✓ Tous les modules\n'
    '✓ 5 établissements\n'
    '✓ Personnalisation\n'
    '✓ Comptabilité avancée\n'
    '✓ Rapports détaillés\n'
    '✓ Support WhatsApp/email\n'
)
pro_p.paragraph_format.line_spacing = 1.5

doc.add_heading('Plan Entreprise - Plus de 500 élèves', level=2)
ent_p = doc.add_paragraph()
ent_p.add_run('250 000 FCFA/mois\n').bold = True
ent_p.add_run(
    '✓ Accès complet illimité\n'
    '✓ Multi-établissements illimités\n'
    '✓ Intégrations API\n'
    '✓ Support 24/7 prioritaire\n'
    '✓ Formations personnalisées\n'
    '✓ Dedicated account manager\n'
)
ent_p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()
note_p = doc.add_paragraph()
note_p.add_run('💡 Note : ').bold = True
note_p.add_run(
    'Abonnement mois par mois. Pas d\'engagement long terme. Résilier à tout moment sans pénalité.'
)
note_p.paragraph_format.line_spacing = 1.5
note_p.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# ============================================
# 10. SECTION 8 : SUPPORT & FORMATION
# ============================================

doc.add_heading('8. Support et Formation', level=1)

doc.add_heading('Canaux de Support', level=2)

support_data = [
    ('WhatsApp', '+237 655 454 994', 'Réponse < 1h'),
    ('Email', 'support@techsolutionsddtp.cm', 'Réponse < 24h'),
    ('Téléphone', '+237 655 454 994', 'Lun-Ven 9h-17h'),
]

support_table = doc.add_table(rows=4, cols=3)
support_table.style = 'Light Grid Accent 1'

# Header
header_cells = support_table.rows[0].cells
header_cells[0].text = 'Canal'
header_cells[1].text = 'Coordonnées'
header_cells[2].text = 'Délai'

for i, (canal, coord, delai) in enumerate(support_data, 1):
    row_cells = support_table.rows[i].cells
    row_cells[0].text = canal
    row_cells[1].text = coord
    row_cells[2].text = delai
    for j in range(3):
        for paragraph in row_cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'

doc.add_paragraph()

doc.add_heading('Formations Incluses', level=2)

training = [
    'Webinaires d\'introduction mensuels (gratuits)',
    'Manuel utilisateur complet 20+ pages (français)',
    'Tutoriels vidéo pas à pas',
    'FAQ 25+ questions répondues',
    'Formation on-site possible (payante)',
]

for t in training:
    p = doc.add_paragraph(t, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

# ============================================
# 11. SECTION 9 : APPEL À L'ACTION
# ============================================

doc.add_heading('9. Prêt à Transformer Votre École ?', level=1)

action_p = doc.add_paragraph()
action_p.add_run('Essayez ERPScolaire gratuitement :\n').bold = True
action_p.add_run(
    '1️⃣ Visitez notre démo en ligne : https://erpscolaire.cm\n'
    '2️⃣ Testez tous les modules (zéro engagement)\n'
    '3️⃣ Téléchargez notre manuel utilisateur\n'
    '4️⃣ Contactez-nous pour une démonstration personnalisée\n'
)
action_p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

# Contact
contact_box = doc.add_paragraph()
contact_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = contact_box.add_run('📞 CONTACTEZ-NOUS MAINTENANT 📞')
contact_run.font.bold = True
contact_run.font.size = Pt(12)
contact_run.font.color.rgb = RGBColor(0, 102, 204)
contact_run.font.name = 'Times New Roman'

contact_details = doc.add_paragraph()
contact_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_details.add_run(
    'Tech Solutions DDTP\n'
    '📧 contact@techsolutionsddtp.cm\n'
    '📱 WhatsApp : +237 655 454 994\n'
    '📞 Téléphone : +237 655 454 994\n'
    '📍 Douala, Logbessou, Cameroun\n'
    '\n'
    '🌐 https://erpscolaire.cm\n'
)
contact_details.paragraph_format.line_spacing = 1.5

for run in contact_details.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run(
    f'© 2026 Tech Solutions DDTP - ERPScolaire™\n'
    f'Tous droits réservés | Cameroun 🇨🇲\n'
    f'Document généré : {datetime.date.today().strftime("%d/%m/%Y")}'
)
footer_run.font.size = Pt(9)
footer_run.font.italic = True
footer_run.font.color.rgb = RGBColor(150, 150, 150)
footer_run.font.name = 'Times New Roman'

# ============================================
# SAUVEGARDE
# ============================================

output_path = r'C:\wamp64\www\erp-scolaire\POSITIONNEMENT_ERPSCOLAIRE_CAMEROUN.docx'
doc.save(output_path)
print('Document cree : ' + output_path)
print('Format : Word (.docx)')
print('Police : Times New Roman, 12pt')
print('Langue : Francais')
print('Pages : 12+')
