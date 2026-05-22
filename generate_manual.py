#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# Create document
doc = Document()

# Set default font to Times New Roman, 12pt
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
title = doc.add_paragraph()
title_run = title.add_run('MANUEL D\'UTILISATION COMPLET\nERPScolaire — Version 1.0')
title_run.font.size = Pt(16)
title_run.font.bold = True
title_run.font.name = 'Times New Roman'
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Date
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run(f'Dernière mise à jour : {datetime.date.today().strftime("%d/%m/%Y")}')
subtitle_run.font.size = Pt(11)
subtitle_run.font.italic = True
subtitle_run.font.name = 'Times New Roman'
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Introduction
doc.add_heading('1. INTRODUCTION', level=1)
p = doc.add_paragraph()
p.add_run('Bienvenue dans ERPScolaire').bold = True
p.add_run(', une solution complète de gestion des établissements scolaires développée pour les écoles primaires, collèges, lycées et universités d\'Afrique francophone.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

p = doc.add_paragraph('Ce logiciel offre une plateforme intégrée permettant de gérer tous les aspects administratifs, pédagogiques et comptables de votre établissement.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

# Accès et Authentification
doc.add_heading('2. ACCÈS ET AUTHENTIFICATION', level=1)

doc.add_heading('2.1 Connexion au système', level=2)
p = doc.add_paragraph('Pour accéder à ERPScolaire :')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

items = [
    'Ouvrez votre navigateur web (Chrome, Firefox, Edge, Safari)',
    'Rendez-vous à l\'adresse : http://localhost/erp-scolaire/public/login',
    'Entrez votre identifiant (username) et mot de passe',
    'Cliquez sur « Connexion »'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.2 Comptes de démonstration', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Rôle'
hdr_cells[1].text = 'Identifiant'
hdr_cells[2].text = 'Mot de passe'

data = [
    ('Super Administrateur', 'superadmin', 'Admin@2025'),
    ('Admin Établissement', 'admin.lma', 'Admin@1234'),
    ('Enseignant', 'prof.toure', 'Prof@1234')
]
for i, row_data in enumerate(data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

# Dashboard
doc.add_heading('3. TABLEAU DE BORD', level=1)
p = doc.add_paragraph('Le tableau de bord est votre point d\'accès central. Il affiche :')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

stats = [
    'Nombre total d\'élèves',
    'Nombre de personnels',
    'Nombre de classes',
    'Graphique des paiements sur 6 mois',
    'Répartition par genre (garçons/filles)',
    'Accès rapide aux modules principaux'
]
for stat in stats:
    doc.add_paragraph(stat, style='List Bullet')

# Gestion des Élèves
doc.add_heading('4. GESTION DES ÉLÈVES', level=1)

doc.add_heading('4.1 Consulter la liste des élèves', level=2)
p = doc.add_paragraph('Accédez au menu « Élèves » pour voir tous les élèves de votre établissement. Vous pouvez :')
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for item in ['Rechercher un élève par nom ou prénom', 'Voir les détails d\'un élève', 'Exporter la liste en Excel', 'Importer des élèves via CSV']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.2 Ajouter un nouvel élève', level=2)
steps = [
    'Cliquez sur le bouton « Ajouter un élève »',
    'Remplissez les informations requises (nom, prénom, date de naissance, genre, etc.)',
    'Assignez l\'élève à une classe',
    'Cliquez sur « Enregistrer »'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('4.3 Importer des élèves (CSV)', level=2)
p = doc.add_paragraph('Pour importer plusieurs élèves à la fois :')
p_format = p.paragraph_format
p_format.line_spacing = 1.5
doc.add_paragraph('Cliquez sur « Importer des élèves »', style='List Bullet')
doc.add_paragraph('Téléchargez le modèle CSV fourni', style='List Bullet')
doc.add_paragraph('Remplissez le fichier avec les données des élèves', style='List Bullet')
doc.add_paragraph('Retournez au formulaire d\'import et téléversez votre fichier', style='List Bullet')
doc.add_paragraph('Vérifiez les données importées', style='List Bullet')

# Gestion du Personnel
doc.add_heading('5. GESTION DU PERSONNEL', level=1)

doc.add_heading('5.1 Consulter le personnel', level=2)
p = doc.add_paragraph('Le module Personnel regroupe tous les employés de l\'établissement : enseignants, administrateurs, secrétaires, comptables, etc.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

doc.add_heading('5.2 Ajouter un membre du personnel', level=2)
steps = [
    'Accédez à « Personnel »',
    'Cliquez sur « Ajouter un personnel »',
    'Entrez les informations (nom, prénom, email, téléphone, fonction)',
    'Assignez un rôle (Enseignant, Admin, Comptable, Secrétaire)',
    'Enregistrez'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

# Classes et Niveaux
doc.add_heading('6. GESTION DES CLASSES', level=1)

doc.add_heading('6.1 Structure des classes', level=2)
p = doc.add_paragraph('Les classes sont organisées par niveaux scolaires. Un établissement peut avoir plusieurs niveaux : Primaire, Collège, Lycée, Université.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

doc.add_heading('6.2 Créer une nouvelle classe', level=2)
steps = [
    'Allez dans « Classes »',
    'Cliquez sur « Ajouter une classe »',
    'Sélectionnez le niveau (2nde, 1ère, Terminale, etc.)',
    'Donnez un identifiant à la classe (ex: « 2nde A »)',
    'Assignez un professeur principal',
    'Enregistrez'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

# Emploi du Temps
doc.add_heading('7. EMPLOI DU TEMPS', level=1)

p = doc.add_paragraph('L\'emploi du temps organise les cours par classe et par semaine.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

doc.add_heading('7.1 Consulter un emploi du temps', level=2)
steps = [
    'Allez dans « Emploi du temps »',
    'Sélectionnez la classe',
    'La grille hebdomadaire s\'affiche avec tous les cours',
    'Identifiez le professeur, la matière et l\'horaire'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

# Présences
doc.add_heading('8. GESTION DES PRÉSENCES', level=1)

doc.add_heading('8.1 Faire l\'appel', level=2)
p = doc.add_paragraph('Les enseignants peuvent faire l\'appel numérique directement dans le système.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

steps = [
    'Cliquez sur « Présences »',
    'Sélectionnez votre cours du jour',
    'Cliquez sur « Faire l\'appel »',
    'Pour chaque élève, marquez : Présent, Absent, Retard ou Excusé',
    'Enregistrez'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('8.2 Consulter les rapports de présence', level=2)
p = doc.add_paragraph('Accédez à « Rapports » → « Présences » pour voir :')
p_format = p.paragraph_format
p_format.line_spacing = 1.5
for item in ['Taux de présence par élève', 'Nombre d\'absences par classe', 'Détails des absences justifiées/non justifiées']:
    doc.add_paragraph(item, style='List Bullet')

# Notes et Bulletins
doc.add_heading('9. SAISIE DES NOTES', level=1)

doc.add_heading('9.1 Saisir des notes', level=2)
steps = [
    'Allez dans « Notes »',
    'Sélectionnez le trimestre (T1, T2, T3)',
    'Choisissez votre classe et matière',
    'Entrez les notes pour chaque élève',
    'Validez'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('9.2 Générer les bulletins', level=2)
p = doc.add_paragraph('Rendez-vous dans « Bulletins » pour générer les relevés de notes.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5
steps = [
    'Sélectionnez le trimestre',
    'Sélectionnez la classe',
    'Le système calcule automatiquement les moyennes',
    'Téléchargez ou imprimez les bulletins en PDF'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

# Paiements
doc.add_heading('10. GESTION DES PAIEMENTS', level=1)

doc.add_heading('10.1 Suivi des paiements', level=2)
p = doc.add_paragraph('Consultez l\'état des paiements pour chaque élève dans le module « Paiements ».')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

doc.add_heading('10.2 États de paiement', level=2)
states = [
    'Soldé : la totalité des frais est payée',
    'Partiel : une partie des frais est payée',
    'Non payé : aucun paiement enregistré'
]
for state in states:
    doc.add_paragraph(state, style='List Bullet')

doc.add_heading('10.3 Enregistrer un paiement', level=2)
steps = [
    'Allez dans « Paiements »',
    'Recherchez le dossier de l\'élève',
    'Cliquez sur « Ajouter un paiement »',
    'Entrez le montant et la date',
    'Sélectionnez le mode de paiement (espèces, chèque, virement)',
    'Enregistrez'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

# Examens
doc.add_heading('11. PROGRAMMATION DES EXAMENS', level=1)

p = doc.add_paragraph('Organisez et tracez tous les examens et compositions.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

steps = [
    'Allez dans « Examens »',
    'Cliquez sur « Créer un examen »',
    'Entrez le nom, type et dates',
    'Associez les classes concernées',
    'Enregistrez'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

# Bibliothèque
doc.add_heading('12. GESTION DE LA BIBLIOTHÈQUE', level=1)

doc.add_heading('12.1 Catalogue des livres', level=2)
p = doc.add_paragraph('Gérez l\'inventaire des livres et les emprunts.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

doc.add_heading('12.2 Ajouter un livre', level=2)
steps = [
    'Cliquez sur « Ajouter un livre »',
    'Entrez titre, auteur, ISBN, éditeur',
    'Spécifiez le nombre d\'exemplaires',
    'Indiquez la localisation (étagère)',
    'Enregistrez'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('12.3 Enregistrer un emprunt', level=2)
steps = [
    'Sélectionnez un livre disponible',
    'Cliquez sur « Prêter »',
    'Choisissez l\'emprunteur',
    'Définissez la date de retour',
    'Validez'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

# Communication
doc.add_heading('13. MODULE DE COMMUNICATION', level=1)

doc.add_heading('13.1 Annonces', level=2)
p = doc.add_paragraph('Publiez des annonces pour informer la communauté scolaire.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

doc.add_heading('13.2 Messagerie interne', level=2)
p = doc.add_paragraph('Envoyez des messages privés aux élèves, enseignants et parents.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

# Rapports
doc.add_heading('14. GÉNÉRATION DE RAPPORTS', level=1)

doc.add_heading('14.1 Rapports disponibles', level=2)
reports = [
    'Rapport des élèves (liste complète avec statuts)',
    'Rapport des notes (moyennes par classe/matière)',
    'Rapport des paiements (synthèse financière)',
    'Rapport des présences (statistiques d\'assiduité)'
]
for report in reports:
    doc.add_paragraph(report, style='List Bullet')

doc.add_heading('14.2 Exporter en Excel', level=2)
p = doc.add_paragraph('Tous les rapports peuvent être exportés en fichier Excel pour analyse dans un tableur.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

# Comptabilité
doc.add_heading('15. GESTION COMPTABLE', level=1)

p = doc.add_paragraph('Suivez les recettes et dépenses de l\'établissement.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

steps = [
    'Allez dans « Comptabilité »',
    'Enregistrez les transactions (recettes/dépenses)',
    'Consultez le bilan financier',
    'Exportez les résumés mensuels/trimestriels'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

# Paramètres
doc.add_heading('16. CONFIGURATION DU SYSTÈME', level=1)

p = doc.add_paragraph('Accédez à « Paramètres » pour configurer votre établissement.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

doc.add_heading('16.1 Paramètres disponibles', level=2)
params = [
    'Année scolaire en cours',
    'Cycles scolaires',
    'Niveaux (Primaire, Collège, Lycée, Université)',
    'Salles de classe',
    'Créneaux horaires',
    'Matières et coefficients'
]
for param in params:
    doc.add_paragraph(param, style='List Bullet')

# Profil Utilisateur
doc.add_heading('17. MON PROFIL', level=1)

p = doc.add_paragraph('Gérez vos informations personnelles.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

doc.add_heading('17.1 Modifications possibles', level=2)
for item in ['Votre photo de profil', 'Votre adresse email', 'Votre mot de passe']:
    doc.add_paragraph(item, style='List Bullet')

# Support
doc.add_heading('18. SUPPORT ET AIDE', level=1)

p = doc.add_paragraph('Pour toute question ou problème technique, contactez notre équipe de support.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

doc.add_heading('18.1 Informations de contact', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Canal'
hdr_cells[1].text = 'Coordonnées'

contact_data = [
    ('Email', 'support@erpscolaire.cm'),
    ('Téléphone', '+237 6XX XXX XXX'),
    ('Site Web', 'www.erpscolaire.cm')
]
for i, (canal, coords) in enumerate(contact_data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = canal
    row_cells[1].text = coords

# FAQ
doc.add_heading('19. QUESTIONS FRÉQUENTES', level=1)

faqs = [
    ('Comment réinitialiser mon mot de passe ?', 'Cliquez sur « Mot de passe oublié » sur la page de connexion et suivez les instructions.'),
    ('Puis-je ajouter plusieurs établissements ?', 'Oui, ERPScolaire supporte plusieurs établissements. Contactez un administrateur.'),
    ('Comment exporter les données ?', 'Utilisez les boutons d\'export disponibles dans chaque module (Excel, PDF).'),
    ('Qui peut voir les notes des élèves ?', 'Les enseignants, directeurs et comptables, selon les permissions configurées.'),
    ('Quelle est la fréquence des sauvegardes ?', 'Les données sont sauvegardées automatiquement à chaque modification.')
]

for question, answer in faqs:
    doc.add_heading(question, level=3)
    p = doc.add_paragraph(answer)
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5

# Conclusion
doc.add_page_break()
doc.add_heading('CONCLUSION', level=1)

p = doc.add_paragraph('ERPScolaire est conçu pour simplifier la gestion quotidienne de votre établissement scolaire. Nous vous remercions de nous faire confiance et nous nous engageons à vous fournir un support de qualité.')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

p = doc.add_paragraph('\nPour toute question supplémentaire ou suggestion d\'amélioration, n\'hésitez pas à nous contacter.\n')
p_format = p.paragraph_format
p_format.line_spacing = 1.5

p = doc.add_paragraph('Bonne utilisation d\'ERPScolaire !')
p_run = p.runs[0]
p_run.bold = True
p_format = p.paragraph_format
p_format.line_spacing = 1.5

# Save document
output_path = r'C:\wamp64\www\erp-scolaire\MANUEL_UTILISATION_ERP_SCOLAIRE_CAMEROUN.docx'
doc.save(output_path)
print(f'Document créé avec succès : {output_path}')
